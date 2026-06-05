from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_EXTRACTED_CHARS = 12000

LOCATION_KEYWORDS = [
    "dubai",
    "abu dhabi",
    "sharjah",
    "ajman",
    "ras al khaimah",
    "rak",
    "fujairah",
    "umm al quwain",
    "al ain",
]

AREA_KEYWORDS = [
    "basement",
    "roof",
    "swimming pool",
    "bathroom",
    "foundation",
    "facade",
    "warehouse",
    "parking",
    "slab",
    "bridge",
    "road",
    "industrial floor",
]

REQUIREMENT_KEYWORDS = [
    "waterproofing",
    "repair",
    "tile adhesive",
    "grout",
    "sealant",
    "coating",
    "crack",
    "corrosion",
    "chemical resistance",
    "floor hardener",
    "curing",
]


class UnsupportedFileType(ValueError):
    pass


def normalize_text(text: str) -> str:
    return " ".join(text.replace("\x00", " ").split())


def trim_text(text: str, limit: int = MAX_EXTRACTED_CHARS) -> str:
    cleaned = normalize_text(text)
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rsplit(" ", 1)[0]


def extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            table_cells.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_cells)


def extract_xlsx(data: bytes) -> str:
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    rows: list[str] = []
    for sheet in workbook.worksheets:
        rows.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(max_row=80, values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                rows.append(" | ".join(values))
    workbook.close()
    return "\n".join(rows)


def extract_text_from_file(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    try:
        if extension == ".txt":
            return trim_text(extract_txt(data))
        if extension == ".pdf":
            return trim_text(extract_pdf(data))
        if extension == ".docx":
            return trim_text(extract_docx(data))
        if extension == ".xlsx":
            return trim_text(extract_xlsx(data))
    except (BadZipFile, OSError, ValueError) as exc:
        raise ValueError("Could not read the uploaded file content.") from exc

    raise UnsupportedFileType("Supported files: PDF, DOCX, XLSX, TXT.")


def find_keywords(text: str, keywords: list[str]) -> list[str]:
    normalized = text.lower()
    return [keyword for keyword in keywords if keyword in normalized]


def summarize_document_signals(text: str) -> dict[str, list[str] | str]:
    preview = trim_text(text, 1200)
    return {
        "preview": preview,
        "locations": find_keywords(text, LOCATION_KEYWORDS),
        "construction_areas": find_keywords(text, AREA_KEYWORDS),
        "requirements": find_keywords(text, REQUIREMENT_KEYWORDS),
    }
